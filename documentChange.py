from docx import Document
from copy import deepcopy

# Load the document
doc = Document(r"path/to/file")

Names = [
"any data",
]


def get_para_data(output_doc_name, paragraph):

    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.style.name = run.style.name

        output_run.font.size=run.font.size
    # Paragraph's alignment data
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
Dictionary = {}
for i in range(len(Names) - 1):
    Dictionary[Names[i]] = Names[i + 1]



for key in Dictionary:
    template = doc.tables[0]
    tbl = template._tbl
    new_tbl = deepcopy(tbl)
    get_para_data(doc, doc.paragraphs[0])
    get_para_data(doc, doc.paragraphs[1])
    get_para_data(doc, doc.paragraphs[2])
    get_para_data(doc, doc.paragraphs[3])
    for row in doc.tables[0].rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    text = run.text
                    if key in text:
                        modified_text = text.replace(key, Dictionary[key])
                        run.text = modified_text

    doc.save(r"path/to/file")

# Add the original content back
    paragraph = doc.add_paragraph()
# After that, we add the previously copied table
    paragraph._p.addnext(new_tbl)

    get_para_data(doc, doc.paragraphs[5])

    doc.save(r"path/to/file")

# Save the modified document



